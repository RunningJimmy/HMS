from docx import Document
from docx.shared import Inches

class WordHandle(object):

    def __init__(self):
        self.doc_obj = Document()

    # 保存
    def close(self,filename):
        self.doc_obj.save(filename)

    # 添加标题
    def add_title(self,title,level= 1):
        '''
        :param title: 标题名
        :param level: 标题等级
        :return:
        '''
        document.add_heading(title,level)

    # 添加分页符
    # 想要下一个文本在一个单独的页面，即使你所在的一个不是满的
    def add_page(self):
        document.add_page_break()


# 电子导检单 HTML 形式
def guide_page_html(session):
    pass

def guide_page_word(session):
    pass


document = Document()



# 添加表格
#默认情况下表格是没有边框的，Table Grid格式是普通的黑色边框表格，更多表格样式可以百度
table = document.add_table(rows=10, cols=4,style="Table Grid")
#自定义表格某一列的宽度，表格的高度。
#默认情况下表格是自动适应文档宽度
#col = table.columns[1]
#设置表格第2列宽度为Inches(5)
#col.width = Inches(5)
for i in range(10):
    for j in range(4):
        cell = table.cell(i, j)
        cell.text = str(i+j)
document.save('C:/Users/Administrator/Desktop/python培训/test.docx')